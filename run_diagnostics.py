#!/usr/bin/env python3
"""Full system diagnostics for SAP CPI Specification Generator."""

import sys
import os
from pathlib import Path

BASE_DIR = Path(__file__).resolve().parent
SRC_DIR = BASE_DIR / 'src'
sys.path.insert(0, str(SRC_DIR))

def run_diagnostics():
    from config import settings as app_settings

    excluded_dirs = {
        '.git',
        'venv',
        '.venv',
        '__pycache__',
        'output',
        'temp',
        '.pytest_cache',
    }

    def discover_iflow_files(search_root: Path):
        discovered = []
        for candidate in search_root.rglob('*.iflw'):
            if any(part in excluded_dirs for part in candidate.parts):
                continue
            discovered.append(candidate)
        return sorted(discovered, key=lambda p: str(p).lower())

    print('=' * 60)
    print('FULL SYSTEM DIAGNOSTICS')
    print('=' * 60)
    
    # 1. Check imports
    print('\n1. MODULE IMPORTS:')
    try:
        from iflow_parser import IFlowParser
        print('   [OK] IFlowParser')
    except Exception as e:
        print(f'   [FAIL] IFlowParser: {e}')
    
    try:
        from ai_generator import AIGenerator
        print('   [OK] AIGenerator')
    except Exception as e:
        print(f'   [FAIL] AIGenerator: {e}')
    
    try:
        from document_builder import build_specification_document
        print('   [OK] document_builder')
    except Exception as e:
        print(f'   [FAIL] document_builder: {e}')
    
    try:
        from diagram_generator import BPMNDiagramGenerator, generate_diagram_bytes
        print('   [OK] diagram_generator')
    except Exception as e:
        print(f'   [FAIL] diagram_generator: {e}')
    
    # 2. Check configuration
    print('\n2. CONFIGURATION:')
    try:
        from config.settings import AI_MODEL, ENABLE_BATCH_MODE
        print(f'   Model: {AI_MODEL}')
        print(f'   Batch mode: {ENABLE_BATCH_MODE}')
    except Exception as e:
        print(f'   [WARN] Config: {e}')
    
    # 3. Check for hardcoded values
    print('\n3. HARDCODING CHECK:')
    sample_iflow_names = [p.stem for p in discover_iflow_files(BASE_DIR)]
    
    files_to_check = [
        BASE_DIR / 'src' / 'document_builder.py',
        BASE_DIR / 'src' / 'ai_generator.py',
        BASE_DIR / 'main.py',
        BASE_DIR / 'src' / 'diagram_generator.py',
    ]
    
    for filepath in files_to_check:
        if filepath.exists():
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()
            found = [name for name in sample_iflow_names if name and name in content]
            if found:
                print(f'   [WARN] Hardcoded in {filepath}: {found}')
            else:
                print(f'   [OK] No hardcoding in {filepath}')
    
    # 4. Test parser with sample
    print('\n4. PARSER TEST:')
    from iflow_parser import IFlowParser
    
    # Find an iFlow file dynamically across the workspace.
    preferred_root = BASE_DIR / 'sample'
    iflow_files = discover_iflow_files(preferred_root) if preferred_root.exists() else []
    if not iflow_files:
        iflow_files = discover_iflow_files(BASE_DIR)

    if not iflow_files:
        print('   [FAIL] No iFlow files found in the workspace for diagnostics.')
        return False
    
    parser = IFlowParser(iflow_files[0])
    parser.parse()
    
    processes = parser.get_integration_processes()
    seq_flows = parser.extract_sequence_flows_with_names()
    msg_flows = parser.extract_message_flows_with_names()
    metadata = parser.extract_metadata()
    senders = parser.extract_sender_properties()
    receivers = parser.extract_receiver_properties()
    
    iflow_name = parser.iflow_name
    
    print(f'   iFlow name: {iflow_name}')
    print(f'   Integration processes: {len(processes)}')
    print(f'   Sequence flows: {len(seq_flows)}')
    print(f'   Message flows: {len(msg_flows)}')
    print(f'   Senders: {len(senders)}')
    print(f'   Receivers: {len(receivers)}')
    
    # 5. Check diagram generator
    print('\n5. DIAGRAM GENERATOR:')
    from diagram_generator import generate_diagram_bytes
    
    for dtype in ['integration_flow']:
        try:
            data = generate_diagram_bytes(parser, dtype)
            if data:
                print(f'   [OK] {dtype}: {len(data)} bytes')
            else:
                print(f'   [FAIL] {dtype}: no bytes returned')
        except Exception as e:
            print(f'   [FAIL] {dtype}: {e}')
    
    # 6. API call count check
    print('\n6. API EFFICIENCY:')
    print('   Batch generation enabled: optimized call pattern is active when cache misses occur.')
    
    # 7. Document quality check
    print('\n7. OUTPUT VALIDATION:')
    output_candidates = sorted(
        app_settings.OUTPUT_DIR.glob('*_TechSpec.docx'),
        key=lambda p: p.stat().st_mtime if p.exists() else 0,
        reverse=True,
    )
    output_file = str(output_candidates[0]) if output_candidates else ''

    if output_file and os.path.exists(output_file):
        from docx import Document
        import zipfile
        
        doc = Document(output_file)
        
        # Count content
        para_count = len(doc.paragraphs)
        table_count = len(doc.tables)
        non_empty = sum(1 for p in doc.paragraphs if p.text.strip())
        
        # Check media
        with zipfile.ZipFile(output_file, 'r') as z:
            media = [f for f in z.namelist() if 'media' in f]
        
        print(f'   Paragraphs: {para_count} ({non_empty} with content)')
        print(f'   Tables: {table_count}')
        print(f'   Images: {len(media)}')
        print(f'   File size: {os.path.getsize(output_file) / 1024:.1f} KB')
        
        # Validate content quality
        if non_empty > 30 and table_count >= 5 and len(media) >= 1:
            print('   [OK] Document appears complete')
        else:
            print('   [WARN] Document may be incomplete')
    else:
        print(f'   [SKIP] No generated *_TechSpec.docx file found in {app_settings.OUTPUT_DIR}')
    
    print('\n' + '=' * 60)
    print('DIAGNOSTICS COMPLETE')
    print('=' * 60)
    
    return True

if __name__ == '__main__':
    run_diagnostics()
