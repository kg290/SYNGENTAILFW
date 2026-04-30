"""
Document Builder - Enterprise-Grade SAP CPI Technical Specification Generator.

Features:
- Professional enterprise document formatting
- BPMN-style diagram integration
- Colored tables with proper spacing
- Clean cover page
- Table of contents
- Proper headers and footers
- Professional typography
- Production-ready output
"""

from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple
from datetime import datetime
import sys
import logging
import io
import re

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

sys.path.insert(0, str(Path(__file__).parent.parent))
from config.settings import (
    OUTPUT_DIR,
    DOC_AUTHOR,
    DOC_VERSION,
    ENABLE_BATCH_MODE,
    TECH_SPEC_SCOPE_MODE,
)

logger = logging.getLogger(__name__)

W15_WORDML_NS = "http://schemas.microsoft.com/office/word/2012/wordml"


class DocumentBuilderError(Exception):
    """Custom exception for document building errors."""
    pass


class EnterpriseDocumentBuilder:
    """Creates enterprise-grade Word documents for SAP CPI technical specs."""
    
    # Professional color palette
    PRIMARY_BLUE = RGBColor(0, 82, 147)
    SECONDARY_BLUE = RGBColor(0, 120, 180)
    DARK_TEXT = RGBColor(33, 37, 41)
    MUTED_TEXT = RGBColor(108, 117, 125)
    TABLE_HEADER_BG = "005293"
    TABLE_ALT_ROW = "F5F8FA"
    
    def __init__(
        self,
        iflow_name: str,
        output_dir: Optional[Path] = None,
        runtime_parameters: Optional[Dict[str, str]] = None,
    ):
        self.iflow_name = iflow_name
        self.output_dir = output_dir or OUTPUT_DIR
        self.output_dir.mkdir(exist_ok=True)
        self.doc = Document()
        self._bookmark_id_counter = 1
        self._runtime_parameters = self._normalize_runtime_parameters(runtime_parameters or {})
        self._ensure_wordml_namespaces()
        self._setup_page()
        self._set_modern_compatibility_mode()
        self._setup_header_footer()
        self._disable_auto_field_updates()

    @staticmethod
    def _normalize_lookup_key(value: str) -> str:
        """Normalize property keys to robust lookup tokens."""
        return re.sub(r'[^a-z0-9]+', '', str(value or '').strip().lower())

    def _normalize_runtime_parameters(self, values: Dict[str, str]) -> Dict[str, str]:
        """Normalize runtime parameter keys for placeholder resolution."""
        normalized: Dict[str, str] = {}
        for key, value in values.items():
            lookup = self._normalize_lookup_key(str(key))
            if lookup and lookup not in normalized:
                normalized[lookup] = self._decode_runtime_value(str(value))
        return normalized

    @staticmethod
    def _decode_runtime_value(value: str) -> str:
        """Decode common Java-properties escape sequences for display."""
        rendered = str(value or '')
        rendered = rendered.replace('\\:', ':')
        rendered = rendered.replace('\\=', '=')
        rendered = rendered.replace('\\ ', ' ')
        return rendered

    def _resolve_runtime_placeholders(self, value: str) -> str:
        """Resolve ${...} and {{...}} placeholders using runtime parameters."""
        raw = str(value or '')
        if not raw:
            return raw

        pattern = re.compile(r'\$\{([^}]+)\}|\{\{([^}]+)\}\}')

        def replacer(match: re.Match[str]) -> str:
            key = (match.group(1) or match.group(2) or '').replace('\\ ', ' ').strip()
            if not key:
                return match.group(0)
            lookup = self._normalize_lookup_key(key)
            if lookup not in self._runtime_parameters:
                return match.group(0)
            return self._runtime_parameters.get(lookup, '')

        return pattern.sub(replacer, self._decode_runtime_value(raw))

    def _ensure_wordml_namespaces(self):
        """Ensure helper namespace map supports w15 element creation."""
        nsmap.setdefault('w15', W15_WORDML_NS)
    
    def _setup_page(self):
        """Configure page layout."""
        section = self.doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1)

    def _set_modern_compatibility_mode(self):
        """Set Word compatibility mode to modern format so collapse behavior is honored."""
        settings = self.doc.settings.element
        compat = settings.find(qn('w:compat'))
        if compat is None:
            compat = OxmlElement('w:compat')
            settings.append(compat)

        for child in list(compat):
            if child.tag == qn('w:compatSetting') and child.get(qn('w:name')) == 'compatibilityMode':
                compat.remove(child)

        compat_setting = OxmlElement('w:compatSetting')
        compat_setting.set(qn('w:name'), 'compatibilityMode')
        compat_setting.set(qn('w:uri'), 'http://schemas.microsoft.com/office/word')
        compat_setting.set(qn('w:val'), '16')
        compat.append(compat_setting)

    def _append_field(self, paragraph, instruction: str, default_text: str = ""):
        """Append a Word field code to a paragraph."""
        run = paragraph.add_run()
        fld_char_begin = OxmlElement('w:fldChar')
        fld_char_begin.set(qn('w:fldCharType'), 'begin')

        instr_text = OxmlElement('w:instrText')
        instr_text.set(qn('xml:space'), 'preserve')
        instr_text.text = instruction

        fld_char_sep = OxmlElement('w:fldChar')
        fld_char_sep.set(qn('w:fldCharType'), 'separate')

        fld_char_end = OxmlElement('w:fldChar')
        fld_char_end.set(qn('w:fldCharType'), 'end')

        run._r.append(fld_char_begin)
        run._r.append(instr_text)
        run._r.append(fld_char_sep)
        if default_text:
            run._r.append(OxmlElement('w:t'))
            run._r[-1].text = default_text
        run._r.append(fld_char_end)

    def _setup_header_footer(self):
        """Add a professional header and page numbers."""
        section = self.doc.sections[0]

        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.text = f"{self.iflow_name} - Technical Specification"
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in header_para.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(9)
            run.font.color.rgb = self.MUTED_TEXT

        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.add_run("Page ")
        self._append_field(footer_para, "PAGE", "1")
        footer_para.add_run(" of ")
        self._append_field(footer_para, "NUMPAGES", "1")

        for run in footer_para.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(9)
            run.font.color.rgb = self.MUTED_TEXT

    def _enable_auto_field_updates(self):
        """Ask Word to update fields (like TOC and page numbers) on open."""
        settings = self.doc.settings.element
        update_fields = settings.find(qn('w:updateFields'))
        if update_fields is None:
            update_fields = OxmlElement('w:updateFields')
            settings.append(update_fields)
        update_fields.set(qn('w:val'), 'true')

    def _disable_auto_field_updates(self):
        """Remove the open-time field update flag to avoid Word prompt on open."""
        settings = self.doc.settings.element
        update_fields = settings.find(qn('w:updateFields'))
        if update_fields is not None:
            settings.remove(update_fields)
    
    def _set_cell_shading(self, cell, color_hex: str):
        """Set background color for a table cell."""
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color_hex)
        cell._tc.get_or_add_tcPr().append(shading)

    def _set_repeat_table_header(self, row):
        """Repeat table header row when a table spans multiple pages."""
        tr_pr = row._tr.get_or_add_trPr()
        tbl_header = OxmlElement('w:tblHeader')
        tbl_header.set(qn('w:val'), 'true')
        tr_pr.append(tbl_header)

    def _set_row_cant_split(self, row):
        """Prevent a table row from being split across pages."""
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement('w:cantSplit')
        tr_pr.append(cant_split)

    @staticmethod
    def _strip_markdown_inline(text: str) -> str:
        """Remove simple markdown formatting artifacts from generated prose."""
        if text is None:
            return ""

        cleaned = str(text)

        # [label](url) -> label
        cleaned = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'\1', cleaned)
        # **bold** / __bold__ -> bold
        cleaned = re.sub(r'\*\*(.*?)\*\*', r'\1', cleaned)
        cleaned = re.sub(r'__(.*?)__', r'\1', cleaned)
        # *italic* / _italic_ -> italic (only around words)
        cleaned = re.sub(r'(?<!\*)\*(?!\*)([^*]+?)(?<!\*)\*(?!\*)', r'\1', cleaned)
        # inline code
        cleaned = cleaned.replace('`', '')

        return cleaned.strip()
    
    def add_cover_page(self):
        """Add professional cover page."""
        # Top spacing
        for _ in range(4):
            self.doc.add_paragraph()
        
        # Main title
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("SAP Integration Suite")
        run.font.size = Pt(32)
        run.font.bold = True
        run.font.name = 'Calibri Light'
        run.font.color.rgb = self.PRIMARY_BLUE
        
        # Subtitle
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Cloud Integration")
        run.font.size = Pt(24)
        run.font.name = 'Calibri Light'
        run.font.color.rgb = self.SECONDARY_BLUE
        
        # Spacing
        for _ in range(2):
            self.doc.add_paragraph()
        
        # Document type
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Technical Specification Document")
        run.font.size = Pt(18)
        run.font.name = 'Calibri'
        run.font.italic = True
        run.font.color.rgb = self.MUTED_TEXT
        
        # Spacing
        for _ in range(2):
            self.doc.add_paragraph()
        
        # iFlow name
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(self.iflow_name)
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.name = 'Calibri'
        run.font.color.rgb = self.DARK_TEXT
        
        # More spacing
        for _ in range(6):
            self.doc.add_paragraph()

        meta_lines = [
            f"Version: {DOC_VERSION}",
            f"Author: {DOC_AUTHOR}",
            f"Date: {datetime.today().strftime('%B %d, %Y')}",
            "Status: Draft",
        ]
        for line in meta_lines:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            run.font.color.rgb = self.MUTED_TEXT
        
        self.doc.add_page_break()
    
    def add_toc_placeholder(self, entries: Optional[List[str]] = None):
        """Add a static table of contents without Word TOC field prompts."""
        h = self.doc.add_heading("Table of Contents", level=1)
        for run in h.runs:
            run.font.color.rgb = self.PRIMARY_BLUE

        toc_entries = entries or []
        if not toc_entries:
            self.add_paragraph("Sections will be listed in the document body.")
            self.doc.add_page_break()
            return

        for raw_entry in toc_entries:
            entry = self._strip_markdown_inline(raw_entry)
            if not entry:
                continue

            level_match = re.match(r"^(\d+)\.(\d+)\b", entry)
            indent_level = 1 if level_match else 0

            para = self.doc.add_paragraph(entry)
            para.paragraph_format.left_indent = Inches(0.28 * indent_level)
            para.paragraph_format.space_after = Pt(4)
            for run in para.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
                run.font.color.rgb = self.DARK_TEXT
        
        self.doc.add_page_break()
    
    def _set_paragraph_outline_level(self, paragraph, level: int):
        """Ensure heading outline level is explicit so Word can collapse the section."""
        p_pr = paragraph._p.get_or_add_pPr()
        outline = p_pr.find(qn('w:outlineLvl'))
        if outline is None:
            outline = OxmlElement('w:outlineLvl')
            p_pr.append(outline)
        outline.set(qn('w:val'), str(max(0, min(8, level - 1))))

    def _set_paragraph_collapsed(self, paragraph, collapsed: bool = True):
        """Mark a heading paragraph as collapsed by default across Word versions."""
        p_pr = paragraph._p.get_or_add_pPr()

        for tag in ('w:collapsed', 'w14:collapsed', 'w15:collapsed'):
            node = p_pr.find(qn(tag))
            if node is not None:
                p_pr.remove(node)

        if not collapsed:
            return

        # Legacy marker used by older Word/OOXML consumers.
        legacy = OxmlElement('w:collapsed')
        legacy.set(qn('w:val'), '1')
        p_pr.append(legacy)

        # Default-collapsed marker used by modern Word (Office 2013+).
        modern = OxmlElement('w15:collapsed')
        p_pr.append(modern)

    def add_heading(self, text: str, level: int = 1, collapsed: bool = False):
        """Add heading with styling."""
        normalized_level = min(level, 4)
        h = self.doc.add_heading(text, level=normalized_level)
        self._set_paragraph_outline_level(h, normalized_level)
        h.paragraph_format.keep_with_next = True
        h.paragraph_format.keep_together = True
        for run in h.runs:
            run.font.name = 'Calibri Light'
            if level == 1:
                run.font.size = Pt(18)
                run.font.color.rgb = self.PRIMARY_BLUE
            elif level == 2:
                run.font.size = Pt(14)
                run.font.color.rgb = self.SECONDARY_BLUE
            else:
                run.font.size = Pt(12)
                run.font.color.rgb = self.DARK_TEXT

        if collapsed:
            self._set_paragraph_collapsed(h, True)

        return h
    
    def add_paragraph(self, text: str, bold: bool = False, italic: bool = False):
        """Add paragraph."""
        p = self.doc.add_paragraph()
        run = p.add_run(self._strip_markdown_inline(self._resolve_runtime_placeholders(text)))
        run.bold = bold
        run.italic = italic
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.color.rgb = self.DARK_TEXT
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.15
        return p

    def _next_bookmark_id(self) -> int:
        """Return the next unique Word bookmark id."""
        next_id = self._bookmark_id_counter
        self._bookmark_id_counter += 1
        return next_id

    def add_bookmark(self, paragraph, bookmark_name: str):
        """Attach a Word bookmark to the provided paragraph."""
        name = str(bookmark_name or "").strip()
        if not name:
            name = f"bookmark_{self._next_bookmark_id()}"

        bookmark_id = str(self._next_bookmark_id())

        bookmark_start = OxmlElement('w:bookmarkStart')
        bookmark_start.set(qn('w:id'), bookmark_id)
        bookmark_start.set(qn('w:name'), name)

        bookmark_end = OxmlElement('w:bookmarkEnd')
        bookmark_end.set(qn('w:id'), bookmark_id)

        paragraph._p.insert(0, bookmark_start)
        paragraph._p.append(bookmark_end)
        return name

    def add_internal_hyperlink(self, paragraph, text: str, anchor: str):
        """Add an internal clickable hyperlink to a bookmark anchor."""
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('w:anchor'), str(anchor))
        hyperlink.set(qn('w:history'), '1')

        run = OxmlElement('w:r')
        run_pr = OxmlElement('w:rPr')

        run_style = OxmlElement('w:rStyle')
        run_style.set(qn('w:val'), 'Hyperlink')
        run_pr.append(run_style)

        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0563C1')
        run_pr.append(color)

        underline = OxmlElement('w:u')
        underline.set(qn('w:val'), 'single')
        run_pr.append(underline)

        run.append(run_pr)

        text_node = OxmlElement('w:t')
        text_node.text = self._strip_markdown_inline(text)
        run.append(text_node)

        hyperlink.append(run)
        paragraph._p.append(hyperlink)
        return paragraph

    def add_internal_link_paragraph(self, text: str, anchor: str):
        """Add a paragraph that links to another place in the same document."""
        p = self.doc.add_paragraph()
        self.add_internal_hyperlink(p, text, anchor)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.15
        return p
    
    def add_bullet_list(self, items: List[str]):
        """Add bullet list."""
        for item in items:
            p = self.doc.add_paragraph(
                self._strip_markdown_inline(self._resolve_runtime_placeholders(item)),
                style='List Bullet',
            )
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(11)

    def add_key_value_bullets(self, rows: List[List[str]]):
        """Render key/value rows as bullet items instead of a table."""
        items: List[str] = []
        for row in rows:
            if not row:
                continue
            key = self._strip_markdown_inline(str(row[0]).strip()) if len(row) >= 1 else ""
            value = self._strip_markdown_inline(str(row[1]).strip()) if len(row) >= 2 else ""
            if key and value:
                items.append(f"{key}: {value}")
            elif key:
                items.append(key)
            elif value:
                items.append(value)
        if items:
            self.add_bullet_list(items)

    def add_numbered_list(self, items: List[str]):
        """Add numbered list."""
        for item in items:
            text = self._strip_markdown_inline(str(item).strip())
            if not text:
                continue
            p = self.doc.add_paragraph(text, style='List Number')
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(11)

    def add_numbered_steps_from_text(self, text: str) -> bool:
        """Split text like '1. ... 2. ...' into a real numbered list."""
        compact = re.sub(r"\s+", " ", self._strip_markdown_inline(str(text))).strip()
        if not compact:
            return False

        matches = re.findall(r'(?:^|\s)(\d{1,2})[\.)]\s+(.*?)(?=(?:\s+\d{1,2}[\.)]\s+)|$)', compact)
        if len(matches) < 2:
            return False

        items = [m[1].strip() for m in matches if m[1].strip()]
        if len(items) < 2:
            return False

        self.add_numbered_list(items)
        return True
    
    def add_table(self, headers: List[str], rows: List[List[str]], caption: Optional[str] = None):
        """Add professional table with header styling."""
        if not headers:
            return

        num_cols = len(headers)
        normalized_rows: List[List[str]] = []
        for row_data in rows:
            if row_data is None:
                continue
            cells = [str(cell).strip() if cell is not None else "" for cell in row_data]
            if len(cells) < num_cols:
                cells.extend([""] * (num_cols - len(cells)))
            elif len(cells) > num_cols:
                cells = cells[:num_cols]

            if any(cells):
                normalized_rows.append(cells)

        if not normalized_rows:
            return
        
        if caption:
            p = self.add_paragraph(f"Table: {caption}", bold=True)
            p.paragraph_format.space_after = Pt(4)
        
        table = self.doc.add_table(rows=1, cols=num_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False

        width_map = {
            2: [Cm(5.5), Cm(10.2)],
            3: [Cm(4.0), Cm(5.0), Cm(6.7)],
            4: [Cm(3.2), Cm(4.0), Cm(4.0), Cm(4.5)],
        }
        column_widths = width_map.get(num_cols)
        if column_widths:
            for idx, width in enumerate(column_widths):
                table.columns[idx].width = width
        
        # Header row
        hdr_row = table.rows[0]
        self._set_repeat_table_header(hdr_row)
        self._set_row_cant_split(hdr_row)
        for i, hdr_text in enumerate(headers):
            cell = hdr_row.cells[i]
            cell.text = self._strip_markdown_inline(self._resolve_runtime_placeholders(hdr_text))
            if column_widths:
                cell.width = column_widths[i]
            self._set_cell_shading(cell, self.TABLE_HEADER_BG)
            for para in cell.paragraphs:
                para.paragraph_format.keep_with_next = True
                for run in para.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.size = Pt(10)
                    run.font.name = 'Calibri'
        
        # Data rows
        for row_idx, row_data in enumerate(normalized_rows):
            row = table.add_row()
            self._set_row_cant_split(row)
            for i, cell_text in enumerate(row_data):
                if i < num_cols:
                    cell = row.cells[i]
                    cell.text = (
                        self._strip_markdown_inline(
                            self._resolve_runtime_placeholders(str(cell_text))
                        )
                        if cell_text else ""
                    )
                    if column_widths:
                        cell.width = column_widths[i]
                    if row_idx % 2 == 1:
                        self._set_cell_shading(cell, self.TABLE_ALT_ROW)
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.size = Pt(10)
                            run.font.name = 'Calibri'
        
        self.doc.add_paragraph()
    
    def add_code_block(self, code: str, language: str = "", max_lines: Optional[int] = 40):
        """Add code block."""
        if language:
            p = self.add_paragraph(f"[{language}]", italic=True)
            p.paragraph_format.space_after = Pt(2)
        
        lines = code.strip().split('\n')
        if max_lines is not None and max_lines > 0 and len(lines) > max_lines:
            code = '\n'.join(lines[:max_lines]) + f"\n\n... [{len(lines) - max_lines} more lines]"
        
        p = self.doc.add_paragraph()
        run = p.add_run(code)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        p.paragraph_format.left_indent = Inches(0.25)
    
    def add_image(self, image_bytes: bytes, width: float = 6.0, caption: Optional[str] = None):
        """Add image from bytes."""
        try:
            img_stream = io.BytesIO(image_bytes)
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(img_stream, width=Inches(width))
            
            if caption:
                cap = self.doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = cap.add_run(f"Figure: {caption}")
                run.font.size = Pt(10)
                run.font.italic = True
                run.font.color.rgb = self.MUTED_TEXT
            
            self.doc.add_paragraph()
        except Exception as e:
            logger.error(f"Failed to add image: {e}")
            self.add_paragraph(f"[Image could not be loaded: {caption}]", italic=True)
    
    def add_page_break(self):
        """Add page break."""
        self.doc.add_page_break()
    
    def save(self, filename: Optional[str] = None) -> Path:
        """Save document."""
        if filename is None:
            safe_name = "".join(c for c in self.iflow_name if c.isalnum() or c in "._- ")
            filename = f"{safe_name}_TechSpec.docx"
        
        output_path = self.output_dir / filename
        self.doc.save(str(output_path))
        logger.info(f"Document saved: {output_path}")
        return output_path


def build_specification_document(
    parser,
    ai_generator,
    groovy_scripts: List[Dict[str, Any]],
    schemas: Optional[List[Dict[str, Any]]] = None,
    parameters: Optional[Dict[str, str]] = None,
    parameter_definitions: Optional[Dict[str, str]] = None,
    all_files: Optional[List[Dict[str, Any]]] = None,
    file_type_summary: Optional[Dict[str, int]] = None,
    text_artifacts: Optional[List[Dict[str, Any]]] = None,
    artifact_analysis_context: str = "",
    output_dir: Optional[Path] = None,
    include_diagrams: bool = True,
    functional_spec_context: str = "",
    functional_spec_analysis: Optional[Dict[str, Any]] = None,
    scope_mode: str = TECH_SPEC_SCOPE_MODE,
) -> Path:
    """Build a detailed enterprise specification document."""

    schemas = schemas or []
    parameters = parameters or {}
    parameter_definitions = parameter_definitions or {}
    all_files = all_files or []
    file_type_summary = file_type_summary or {}
    text_artifacts = text_artifacts or []
    functional_spec_analysis = functional_spec_analysis or {}
    scope_mode = str(scope_mode or TECH_SPEC_SCOPE_MODE).strip().lower()
    is_extended_scope = scope_mode == "extended"

    iflow_name = parser.iflow_name
    all_processes = parser.get_integration_processes()
    main_processes = all_processes[:1]
    # Only additional top-level BPMN processes are treated as local integration processes.
    # Exception subprocesses stay in the dedicated error-handling section.
    local_processes = all_processes[1:]
    message_flows = parser.extract_message_flows_with_names()
    sequence_flows = parser.extract_sequence_flows_with_names()
    sender_props = parser.extract_sender_properties()
    receiver_props = parser.extract_receiver_properties()
    mapping_props = parser.extract_mapping_properties()
    security_props = parser.extract_security_properties()
    exception_props = parser.extract_exception_properties()
    metadata = parser.extract_metadata()

    builder = EnterpriseDocumentBuilder(
        iflow_name,
        output_dir,
        runtime_parameters=parameters,
    )

    print("\n" + "=" * 70)
    print("  SAP CI Technical Specification Generator")
    print("=" * 70)
    print(f"  iFlow: {iflow_name}")
    print(f"  Processes: {len(all_processes)}")
    print(f"  Scripts: {len(groovy_scripts)}")
    print(f"  Schemas: {len(schemas)}")
    print(f"  Total Files Analyzed: {len(all_files)}")
    print(f"  Diagrams: {'Enabled' if include_diagrams else 'Disabled'}")
    print("=" * 70)

    # ========================================================================
    # GENERATE DIAGRAMS
    # ========================================================================
    diagram_bytes: Dict[str, bytes] = {}
    local_process_diagrams: List[Tuple[str, bytes]] = []
    if include_diagrams:
        print("\n[INFO] Generating diagrams...")
        try:
            from src.diagram_generator import (
                generate_diagram_bytes,
                extract_exception_subdiagram_bytes,
                generate_process_diagram_bytes,
            )

            for dtype in ['integration_flow', 'sender', 'receiver']:
                print(f"   - {dtype.replace('_', ' ').title()}...", end=" ")
                try:
                    img = generate_diagram_bytes(parser, dtype)
                    if img:
                        diagram_bytes[dtype] = img
                        print("[OK]")
                    else:
                        print("[WARN]")
                except Exception as e:
                    print(f"[ERROR] {e}")

            if 'integration_flow' in diagram_bytes:
                print("   - Exception Subprocess Snapshot...", end=" ")
                try:
                    exception_img = extract_exception_subdiagram_bytes(
                        parser,
                        diagram_bytes['integration_flow'],
                    )
                    if exception_img:
                        diagram_bytes['exception_subprocess'] = exception_img
                        print("[OK]")
                    else:
                        print("[SKIP]")
                except Exception as e:
                    print(f"[WARN] {e}")

            if local_processes:
                for idx, process in enumerate(local_processes, start=1):
                    process_name = str(process.get("name") or f"Local Process {idx}")
                    print(f"   - Local Process Diagram ({process_name})...", end=" ")
                    try:
                        local_img = generate_process_diagram_bytes(parser, process)
                        if local_img:
                            local_process_diagrams.append((process_name, local_img))
                            print("[OK]")
                        else:
                            print("[SKIP]")
                    except Exception as e:
                        print(f"[WARN] {e}")
        except ImportError as e:
            print(f"   [WARN] Diagrams not available: {e}")

    # ========================================================================
    # BATCH AI GENERATION
    # ========================================================================
    batch: Dict[str, Any] = {}

    if ENABLE_BATCH_MODE and hasattr(ai_generator, 'generate_all_sections_batch'):
        print("\n[INFO] Generating AI content...")

        combined_xml = "\n".join([
            parser.extract_collaboration_xml(),
            parser.extract_process_xml(),
            parser.extract_message_flows_xml(),
            parser.sender_props_to_xml(),
            parser.receiver_props_to_xml(),
            parser.mapping_props_to_xml(),
            parser.exception_props_to_xml(),
            parser.metadata_to_xml(),
        ])

        generated = ai_generator.generate_all_sections_batch(
            iflow_name=iflow_name,
            xml_content=combined_xml,
            groovy_scripts=groovy_scripts,
            functional_spec_context=functional_spec_context,
            functional_spec_analysis=functional_spec_analysis,
            artifact_analysis_context=artifact_analysis_context,
        )

        if isinstance(generated, dict):
            batch = generated
            print("   [OK] AI content generated")
        else:
            print("   [WARN] Using fallback mode")

    def value_to_text(value: Any) -> str:
        if value is None:
            return ""
        if isinstance(value, list):
            return ", ".join(str(v) for v in value if str(v).strip())
        if isinstance(value, dict):
            parts = []
            for k, v in value.items():
                rendered = value_to_text(v)
                if rendered:
                    parts.append(f"{k.replace('_', ' ').title()}: {rendered}")
            return "; ".join(parts)
        return str(value)

    def get_text(key: str, default: str = "") -> str:
        if key not in batch:
            return default
        rendered = value_to_text(batch.get(key))
        return rendered if rendered else default

    def get_dict(key: str) -> Dict[str, Any]:
        value = batch.get(key, {})
        return value if isinstance(value, dict) else {}

    def get_list(key: str) -> List[Any]:
        value = batch.get(key, [])
        return value if isinstance(value, list) else []

    def dict_to_bullets(data: Dict[str, Any]) -> List[str]:
        items: List[str] = []
        for k, v in data.items():
            rendered = value_to_text(v)
            if rendered:
                items.append(f"{k.replace('_', ' ').title()}: {rendered}")
        return items

    ai_stats: Dict[str, Any] = {}
    if hasattr(ai_generator, "get_stats"):
        raw_stats = ai_generator.get_stats()
        if isinstance(raw_stats, dict):
            ai_stats = raw_stats

    assumptions = get_dict("functional_assumptions")
    functional_alignment = get_dict("functional_spec_alignment")
    artifact_coverage = get_dict("artifact_coverage")
    dependencies_text = get_text("technical_dependencies")
    process_flow_data = get_dict("process_flow")
    include_traceability_section = is_extended_scope and bool(
        functional_spec_analysis or functional_alignment
    )

    def dict_to_rows(data: Dict[str, Any]) -> List[List[str]]:
        rows: List[List[str]] = []
        for k, v in data.items():
            rendered = value_to_text(v)
            if rendered:
                rows.append([k.replace('_', ' ').title(), rendered])
        return rows

    def props_to_map(rows: List[List[str]]) -> Dict[str, str]:
        normalized: Dict[str, str] = {}
        for key, value in rows:
            lookup = builder._normalize_lookup_key(str(key))
            rendered = builder._resolve_runtime_placeholders(str(value))
            if lookup and rendered and lookup not in normalized:
                normalized[lookup] = rendered
        return normalized

    def list_artifact_paths(extensions: List[str]) -> List[str]:
        allowed = {ext.lower() for ext in extensions}
        matches: List[str] = []
        for entry in all_files:
            ext = str(entry.get("extension", "")).lower()
            rel_path = str(entry.get("relative_path", "")).strip()
            if rel_path and ext in allowed:
                matches.append(rel_path)
        return matches

    def label_from_identifier(text: str) -> str:
        base = Path(str(text or "")).stem or str(text or "")
        parts = [part for part in re.split(r'[_\-/]+', base) if part]
        if not parts:
            return str(text or "").strip()
        return " ".join(parts)

    def tidy_mapping_label(text: str) -> str:
        raw = str(text or "").strip()
        if not raw:
            return "Not identified"
        cleaned = re.sub(r'^(mm|xslt|map)\s+', '', label_from_identifier(raw), flags=re.IGNORECASE).strip()
        cleaned = re.sub(r'\b(xslt|mapping|mmap)\b', '', cleaned, flags=re.IGNORECASE).strip()
        cleaned = re.sub(r'\s{2,}', ' ', cleaned).strip(" -_")
        return cleaned or "Not identified"

    def pick_prop_value(prop_map: Dict[str, str], candidate_keys: List[str], fallback: str = "") -> str:
        for key in candidate_keys:
            if key == "_protocol":
                rendered = build_protocol_value(prop_map).strip()
            else:
                rendered = prop_map.get(builder._normalize_lookup_key(key), "").strip()
            if rendered:
                return rendered
        return fallback

    def build_protocol_value(prop_map: Dict[str, str]) -> str:
        transport = prop_map.get("transportprotocol", "").strip()
        message = prop_map.get("messageprotocol", "").strip()
        if not transport:
            transport = prop_map.get("transportprotocolversion", "").strip()
        if message.lower() == "none":
            message = ""
        if message and transport:
            return f"{message} over {transport}"
        return message or transport

    def extract_mapping_relation(mapping_name: str, mapping_path: str) -> Tuple[str, str]:
        identifier = str(mapping_name or Path(mapping_path).stem or "").strip()
        raw_parts = [part for part in re.split(r'[_\-/]+', identifier) if part]
        lower_parts = [part.lower() for part in raw_parts]
        if lower_parts and lower_parts[0] in {"mm", "xslt", "map"}:
            raw_parts = raw_parts[1:]
            lower_parts = lower_parts[1:]
        if "to" in lower_parts:
            split_at = lower_parts.index("to")
            source_label = tidy_mapping_label(" ".join(raw_parts[:split_at]))
            target_label = tidy_mapping_label(" ".join(raw_parts[split_at + 1:]))
            return source_label, target_label
        readable = tidy_mapping_label(identifier)
        return readable or "Not identified", "Not identified"

    def build_mapping_summary(mapping_name: str, mapping_path: str) -> List[str]:
        identifier = mapping_name or Path(mapping_path).stem
        readable = label_from_identifier(identifier)
        source_label, target_label = extract_mapping_relation(str(mapping_name), str(mapping_path))

        if target_label != "Not identified":
            line_one = (
                f"This mapping transforms {source_label} data into the {target_label} structure used later in SAP CI."
            )
        else:
            line_one = f"This mapping applies the configured transformation logic for {readable} within SAP CI."

        line_two = (
            "It is executed during message processing to reshape the payload before the next handoff."
        )
        return [line_one, line_two]

    def filter_display_rows(rows: List[List[str]]) -> List[List[str]]:
        filtered: List[List[str]] = []
        skip_keys = {
            "headertable",
            "propertytable",
            "cmdvarianturi",
            "componentswcvname",
            "componentswcvid",
            "componentns",
            "componentversion",
            "activitytype",
            "exprtype",
            "streaming",
            "stoponexecution",
            "splitterthreads",
            "parallelprocessing",
            "grouping",
            "splittype",
            "timeout",
        }
        for key, value in rows:
            key_text = str(key or "").strip()
            value_text = builder._resolve_runtime_placeholders(str(value or "")).strip()
            normalized_key = builder._normalize_lookup_key(key_text)
            if not key_text or not value_text:
                continue
            if normalized_key in skip_keys:
                continue
            if "<row>" in value_text and "<cell" in value_text:
                continue
            filtered.append([key_text, value_text])
        return filtered

    def fallback_process_description(process: Dict[str, Any]) -> str:
        process_elem = process.get("element")
        if process_elem is None:
            return ""

        sequence_count = len(parser.extract_sequence_flows_for_process(process_elem))
        child_names: List[str] = []
        for child in parser.extract_child_properties(process_elem):
            child_tag = str(child.get("tag", "")).strip().lower()
            name = str(child.get("name", "")).strip() or str(child.get("heading", "")).strip()
            if child_tag in {"sequenceflow", "startevent", "endevent", "extensionelements"}:
                continue
            if name and name not in child_names:
                child_names.append(name)

        summary = f"This process contains {sequence_count} sequence flow(s)."
        if child_names:
            summary += " Key activities include " + ", ".join(child_names[:4]) + "."
        return summary

    def fallback_process_steps(process_elem: Any) -> List[str]:
        if process_elem is None:
            return []
        ordered_nodes: List[str] = []
        for src, tgt, _ in parser.extract_sequence_flows_for_process(process_elem):
            for node in (src, tgt):
                node_text = str(node).strip()
                node_lower = node_text.lower()
                if not node_text or "start" in node_lower or "end" in node_lower:
                    continue
                if node_text not in ordered_nodes:
                    ordered_nodes.append(node_text)
        return [f"{node} is executed within the process flow." for node in ordered_nodes[:6]]

    def build_adapter_description(prop_map: Dict[str, str], direction: str) -> str:
        system = prop_map.get("system", "") or ("sender system" if direction == "sender" else "receiver system")
        adapter_type = prop_map.get("componenttype", "") or prop_map.get("name", "") or "configured"
        protocol = build_protocol_value(prop_map)
        if direction == "sender":
            address = pick_prop_value(prop_map, ["address", "urlpath", "addressinbound"])
        else:
            address = pick_prop_value(prop_map, ["address", "soapwsdlurl", "url", "urlpath", "host", "alias"])
        if direction == "sender":
            parts = [f"The {adapter_type} sender adapter receives messages from {system}."]
            if address:
                parts.append(f"It listens on {address}.")
            if protocol:
                parts.append(f"It uses {protocol} for inbound communication.")
            return " ".join(parts)

        auth = prop_map.get("authentication", "")
        credential = pick_prop_value(prop_map, ["credentialname", "usernametokencredentialname", "alias", "accesskey", "secretkey"])
        parts = [f"The {adapter_type} receiver adapter sends processed messages to {system}."]
        if address:
            parts.append(f"It targets {address}.")
        if auth:
            parts.append(f"Authentication type is {auth}.")
        if credential:
            parts.append(f"It uses credential/security artifact {credential}.")
        return " ".join(parts)

    def summarize_processes(processes: List[Dict[str, Any]]) -> List[List[str]]:
        rows: List[List[str]] = []
        for process in processes:
            process_elem = process.get("element")
            if process_elem is None:
                continue

            sequence_count = len(parser.extract_sequence_flows_for_process(process_elem))
            child_count = len(
                [
                    child
                    for child in list(process_elem)
                    if child.tag.split("}")[-1] not in {"sequenceFlow", "extensionElements"}
                ]
            )
            rows.append(
                [
                    str(process.get("name", "Process")),
                    str(sequence_count),
                    str(child_count),
                ]
            )
        return rows

    def render_externalized_parameter_rows(values: Dict[str, str]) -> List[List[str]]:
        rows: List[List[str]] = []
        for key, value in values.items():
            rows.append([str(key), builder._resolve_runtime_placeholders(str(value))])
        return rows

    def parameter_placeholder_names(raw_value: str) -> List[str]:
        matches = re.findall(r'\$\{([^}]+)\}|\{\{([^}]+)\}\}', str(raw_value or ''))
        names: List[str] = []
        for left, right in matches:
            key = (left or right or "").replace('\\ ', ' ').strip()
            if key:
                names.append(key)
        return names

    def build_parameter_usage_rows(values: Dict[str, str]) -> List[List[str]]:
        raw_contexts: List[Tuple[str, str]] = []

        for key, value in sender_props:
            raw_contexts.append((f"Sender - {key}", str(value)))
        for key, value in receiver_props:
            raw_contexts.append((f"Receiver - {key}", str(value)))
        for mapping_idx, mapping in enumerate(mapping_props, start=1):
            mapping_map = {builder._normalize_lookup_key(str(k)): str(v or "") for k, v in mapping}
            mapping_name = (
                mapping_map.get("mappingname")
                or Path(mapping_map.get("mappinguri", "")).stem
                or Path(mapping_map.get("mappingpath", "")).stem
                or f"Mapping {mapping_idx}"
            )
            for key, value in mapping:
                raw_contexts.append((f"Mapping {mapping_name} - {key}", str(value)))

        for process in all_processes:
            process_name = str(process.get("name") or "Process")
            process_elem = process.get("element")
            if process_elem is None:
                continue
            for _, key, value in parser.extract_components_from_process(process_elem):
                raw_contexts.append((f"{process_name} - {key}", str(value)))
            for child in parser.extract_child_properties(process_elem):
                child_heading = str(child.get("heading", "")).strip() or process_name
                for key, value in child.get("properties", []):
                    raw_contexts.append((f"{process_name} / {child_heading} - {key}", str(value)))

        rows: List[List[str]] = []
        for key, value in values.items():
            rendered_value = builder._resolve_runtime_placeholders(str(value))
            lookup = builder._normalize_lookup_key(str(key))
            usages: List[str] = []
            seen_usages: set[str] = set()
            for context_label, raw_value in raw_contexts:
                for placeholder in parameter_placeholder_names(raw_value):
                    if builder._normalize_lookup_key(placeholder) != lookup:
                        continue
                    if context_label in seen_usages:
                        continue
                    seen_usages.add(context_label)
                    usages.append(context_label)
            if not usages:
                continue
            rows.append([
                str(key),
                rendered_value or "Not configured",
                "; ".join(usages),
            ])
        rows.sort(key=lambda item: item[0].lower())
        return rows

    def build_adapter_summary_rows(prop_map: Dict[str, str], direction: str) -> List[List[str]]:
        if direction == "sender":
            ordered_fields = [
                ("Path / Address", ["address", "urlpath", "addressinbound"]),
                ("Adapter Type", ["componenttype", "name"]),
                ("Protocol", ["_protocol"]),
                ("Auth Type", ["senderauthtype", "authorization", "authentication"]),
                ("User Role", ["userrole"]),
            ]
        else:
            ordered_fields = [
                ("URL / Address", ["address", "soapwsdlurl", "url", "urlpath", "host", "alias"]),
                ("Adapter Type", ["componenttype", "name"]),
                ("Authentication Type", ["authentication"]),
                ("Credential / Security Artifact", ["credentialname", "usernametokencredentialname", "alias", "accesskey", "secretkey"]),
                ("Operation Name", ["operationname", "operation", "soapservicename", "s3receiveroperation"]),
            ]

        rows: List[List[str]] = []
        for label, keys in ordered_fields:
            value = pick_prop_value(prop_map, keys, "Not configured")
            rows.append([label, value])

        return rows

    print("\n[INFO] Building document...")

    # Cover + TOC
    print("   [1/10] Cover Page...")
    builder.add_cover_page()
    print("   [2/10] Table of Contents...")
    toc_entries: List[str] = [
        "1. Change History",
        "1.1 Document Control",
        "2. Overview",
        "2.1 Executive Summary",
        "2.2 Purpose",
        "2.3 Interface Requirement",
    ]
    if assumptions:
        toc_entries.append("2.4 Functional Assumptions")
    toc_entries.extend([
        "2.5 Integration Snapshot",
        "3. High level iFlow Design",
        "3.1 Current Scenario",
        "3.2 To-Be Scenario",
        "3.3 High Level Design",
    ])
    if include_traceability_section:
        toc_entries.insert(toc_entries.index("3. High level iFlow Design"), "2.6 Functional Specification Traceability")
    toc_entries.append("3.4 Technical Dependencies")
    toc_entries.extend([
        "4. Message Flow",
        "4.1 Process Flow",
    ])
    if message_flows:
        toc_entries.append("4.2 Message Flow Connections")
    elif sequence_flows:
        toc_entries.append("4.2 Sequence Flow Connections")
    if 'integration_flow' in diagram_bytes:
        toc_entries.append("4.3 Integration Flow Diagram")
    toc_entries.extend([
        "5. Technical Description",
        "5.1 Main Integration Process",
    ])
    if local_processes:
        toc_entries.append("5.2 Local Integration Process")
    toc_entries.extend([
        "5.3 Sender",
        "5.4 Receiver",
        "5.5 Mappings",
        "5.6 Security",
        "5.7 Groovy Scripts",
        "5.8 Error Handling & Logging",
        "6. Version and Metadata",
        "7. Appendix",
        "7.1 Technical Artifacts",
        "7.2 Schemas",
        "7.3 Externalized Parameters",
    ])
    if is_extended_scope:
        toc_entries.extend([
            "7.4 Glossary",
            "7.5 References",
            "7.6 Generation Statistics",
            "7.7 File Type Summary",
            "7.8 File Inventory",
            "7.9 Artifact Evidence",
        ])
    builder.add_toc_placeholder(toc_entries)

    # 1. Change History
    print("   [3/10] Change History...")
    builder.add_heading("1. Change History", 1)
    builder.add_table(
        ["Version", "Date", "Author", "Description"],
        [[DOC_VERSION, datetime.today().strftime("%Y-%m-%d"), DOC_AUTHOR, "Initial autogenerated version"]],
    )

    builder.add_heading("1.1 Document Control", 2)
    builder.add_bullet_list(
        [
            f"Document ID: TS-{iflow_name}",
            f"Integration Flow: {iflow_name}",
            f"Generated On: {datetime.today().strftime('%Y-%m-%d %H:%M')}",
            f"Generation Mode: {'Batch' if ENABLE_BATCH_MODE else 'Section-by-section'}",
            f"Diagrams Included: {'Yes' if include_diagrams else 'No'}",
            f"Functional Spec Context: {'Provided' if functional_spec_context.strip() else 'Not Provided'}",
        ]
    )

    # 2. Overview
    print("   [4/10] Overview...")
    builder.add_page_break()
    builder.add_heading("2. Overview", 1)
    builder.add_heading("2.1 Executive Summary", 2)
    builder.add_paragraph(get_text("executive_summary", f"Technical specification for integration flow {iflow_name}."))

    builder.add_heading("2.2 Purpose", 2)
    builder.add_paragraph(get_text("purpose", "This document describes design and implementation details for the integration flow."))

    builder.add_heading("2.3 Interface Requirement", 2)
    builder.add_paragraph(get_text("interface_requirement", "Business requirement is captured through source-to-target system integration."))

    if assumptions:
        builder.add_heading("2.4 Functional Assumptions", 2)
        builder.add_table(["Assumption", "Value"], dict_to_rows(assumptions))

    builder.add_heading("2.5 Integration Snapshot", 2)
    snapshot_rows = [
        ["Integration Processes", str(len(all_processes))],
        ["Message Flows", str(len(message_flows))],
        ["Sequence Flows", str(len(sequence_flows))],
        ["Groovy Scripts", str(len(groovy_scripts))],
        ["Schemas", str(len(schemas))],
        ["Externalized Parameters", str(len(parameters))],
    ]
    if is_extended_scope:
        snapshot_rows.append(["Total Files Analyzed", str(len(all_files))])

    builder.add_table(
        ["Metric", "Value"],
        snapshot_rows,
    )

    if include_traceability_section:
        builder.add_heading("2.6 Functional Specification Traceability", 2)
        if functional_spec_analysis:
            analysis_summary_rows = dict_to_rows(
                {
                    "analysis_summary": functional_spec_analysis.get("analysis_summary", ""),
                    "document_count": functional_spec_analysis.get("document_count", ""),
                    "source_files": functional_spec_analysis.get("source_files", []),
                    "top_terms": functional_spec_analysis.get("top_terms", []),
                }
            )
            if analysis_summary_rows:
                builder.add_table(["Aspect", "Details"], analysis_summary_rows)

            fs_req = functional_spec_analysis.get("business_requirements", [])
            if isinstance(fs_req, list) and fs_req:
                builder.add_heading("Detected Requirement Signals", 3)
                builder.add_bullet_list([str(item) for item in fs_req])

            fs_interfaces = functional_spec_analysis.get("interface_points", [])
            if isinstance(fs_interfaces, list) and fs_interfaces:
                builder.add_heading("Detected Interface Signals", 3)
                builder.add_bullet_list([str(item) for item in fs_interfaces])

        if functional_alignment:
            builder.add_heading("AI Traceability Interpretation", 3)
            alignment_rows = dict_to_rows(
                {
                    "requirement_traceability": functional_alignment.get("requirement_traceability", []),
                    "assumptions_used": functional_alignment.get("assumptions_used", []),
                    "open_questions": functional_alignment.get("open_questions", []),
                }
            )
            if alignment_rows:
                builder.add_table(["Traceability Aspect", "Details"], alignment_rows)

        if not functional_spec_analysis and not functional_alignment:
            builder.add_paragraph(
                "Functional specification traceability is not available because no readable functional specification evidence was loaded.",
                italic=True,
            )

    # 3. High Level iFlow Design
    print("   [5/10] High Level Design...")
    builder.add_page_break()
    builder.add_heading("3. High level iFlow Design", 1)
    builder.add_heading("3.1 Current Scenario", 2)
    builder.add_paragraph(get_text("current_scenario", "Current scenario details are not explicitly configured in the iFlow."))

    builder.add_heading("3.2 To-Be Scenario", 2)
    builder.add_paragraph(get_text("tobe_scenario", "The integration automates message processing between source and target landscapes."))

    builder.add_heading("3.3 High Level Design", 2)
    builder.add_paragraph(get_text("high_level_design", "The integration is implemented using SAP CI with adapter-based communication."))

    sender_prop_map = props_to_map(sender_props)
    receiver_prop_map = props_to_map(receiver_props)

    dependency_items: List[str] = []
    source_system = (
        builder._resolve_runtime_placeholders(str(process_flow_data.get("source_system", ""))).strip()
        or sender_prop_map.get("system", "")
        or "source system"
    )
    target_system = (
        builder._resolve_runtime_placeholders(str(process_flow_data.get("target_system", ""))).strip()
        or receiver_prop_map.get("system", "")
        or "receiver system"
    )
    dependency_items.append(
        f"Runtime platform: SAP CI integration flow '{iflow_name}' orchestrates message exchange from {source_system} to {target_system}."
    )

    sender_address = pick_prop_value(sender_prop_map, ["address", "urlpath", "addressinbound"])
    sender_protocol = sender_prop_map.get("transportprotocol", "")
    sender_message_protocol = sender_prop_map.get("messageprotocol", "")
    sender_auth = pick_prop_value(sender_prop_map, ["senderauthtype", "authorization", "authentication"])
    sender_role = pick_prop_value(sender_prop_map, ["userrole"])
    inbound_parts = []
    if sender_address:
        inbound_parts.append(f"path {sender_address}")
    if sender_protocol or sender_message_protocol:
        inbound_parts.append(
            f"{sender_prop_map.get('componenttype', 'Not configured')} adapter using "
            f"{sender_message_protocol or 'Not configured'} over {sender_protocol or 'Not configured'}"
        )
    if sender_auth or sender_role:
        inbound_parts.append(
            f"sender auth type {sender_auth or 'Not configured'}"
            + (f" with role {sender_role}" if sender_role else "")
        )
    if inbound_parts:
        dependency_items.append("Inbound interface dependency: " + "; ".join(inbound_parts) + ".")

    receiver_address = pick_prop_value(receiver_prop_map, ["address", "soapwsdlurl", "url", "urlpath", "host", "alias"])

    receiver_auth = pick_prop_value(receiver_prop_map, ["authentication"])
    receiver_credential = pick_prop_value(receiver_prop_map, ["credentialname", "usernametokencredentialname", "alias", "accesskey", "secretkey"])
    fallback_credential = builder._resolve_runtime_placeholders(str(parameters.get("Credential", ""))).strip()
    if not receiver_credential and fallback_credential:
        receiver_credential = fallback_credential
    receiver_operation = pick_prop_value(receiver_prop_map, ["operationname", "operation", "soapservicename", "s3receiveroperation"]) or builder._resolve_runtime_placeholders(
        str(parameters.get("Operation", ""))
    ).strip()
    receiver_processing = receiver_prop_map.get("processing", "") or builder._resolve_runtime_placeholders(
        str(parameters.get("Processing", ""))
    ).strip()
    receiver_ws_security = receiver_prop_map.get("wssecuritytypeoutbound", "")
    receiver_signature = receiver_prop_map.get("wssecuritysignaturealgorithm", "")
    receiver_cert = receiver_prop_map.get("recipientx509tokenassertion", "") or receiver_prop_map.get("x509tokenassertion", "")
    proxy_type = receiver_prop_map.get("proxytype", "")
    proxy_host = receiver_prop_map.get("proxyhost", "")
    proxy_port = receiver_prop_map.get("proxyport", "")
    timeout = receiver_prop_map.get("requesttimeout", "")
    keep_alive = receiver_prop_map.get("keepconnectionalive", "")
    outbound_parts = []
    if receiver_address:
        outbound_parts.append(f"URL {receiver_address}")
    if receiver_auth or receiver_credential:
        outbound_parts.append(
            f"receiver authentication type {receiver_auth or 'Not configured'}"
            + (f" with credential/security artifact {receiver_credential}" if receiver_credential else "")
        )
    if receiver_operation:
        outbound_parts.append(f"operation {receiver_operation}")
    if receiver_processing:
        outbound_parts.append(f"processing mode {receiver_processing}")
    if outbound_parts:
        dependency_items.append("Outbound interface dependency: " + "; ".join(outbound_parts) + ".")

    if sender_auth or sender_role or receiver_auth or receiver_credential:
        security_dependency_parts: List[str] = []
        if sender_auth or sender_role:
            security_dependency_parts.append(
                f"sender auth type {sender_auth or 'Not configured'}"
                + (f" with user role {sender_role}" if sender_role else "")
            )
        if receiver_auth or receiver_credential:
            security_dependency_parts.append(
                f"receiver authentication type {receiver_auth or 'Not configured'}"
                + (f" with credential/security artifact {receiver_credential}" if receiver_credential else "")
            )
        dependency_items.append("Security dependency: " + "; ".join(security_dependency_parts) + ".")

    runtime_security_parts = []
    if receiver_ws_security:
        runtime_security_parts.append(f"WS-Security mode {receiver_ws_security}")
    if receiver_signature:
        runtime_security_parts.append(f"signature algorithm {receiver_signature}")
    if receiver_cert:
        runtime_security_parts.append(f"token profile {receiver_cert}")
    proxy_parts = [part for part in [proxy_type, proxy_host, proxy_port] if part]
    if proxy_parts:
        runtime_security_parts.append(f"proxy {' / '.join(proxy_parts)}")
    if timeout or keep_alive:
        runtime_security_parts.append(
            f"timeout {timeout or 'Not configured'} ms"
            + (f" and keep-alive {keep_alive}" if keep_alive else "")
        )
    if runtime_security_parts:
        dependency_items.append("Runtime security and network dependency: " + "; ".join(runtime_security_parts) + ".")

    wsdl_paths = list_artifact_paths([".wsdl"])
    mapping_paths = list_artifact_paths([".xsl", ".xslt", ".mmap"])
    artifact_parts = []
    if wsdl_paths:
        artifact_parts.append("WSDLs " + "; ".join(wsdl_paths[:2]))
    if mapping_paths:
        artifact_parts.append("mappings " + "; ".join(mapping_paths[:2]))
    if artifact_parts:
        dependency_items.append("Design-time artifacts: " + " | ".join(artifact_parts) + ".")

    groovy_paths = [Path(str(script.get("file_path"))).name for script in groovy_scripts if str(script.get("file_path", "")).strip()]
    if groovy_paths:
        dependency_items.append("Custom script artifacts: " + "; ".join(groovy_paths[:4]) + ".")

    if parameters:
        priority_keys = ["Address", "Address_Mic", "Credential", "Auth", "UserRole", "BusinessObject"]
        highlighted_params = [
            key
            for key in priority_keys
            if builder._resolve_runtime_placeholders(str(parameters.get(key, ""))).strip()
        ]
        dependency_items.append(
            f"Externalized runtime dependencies: {len(parameters)} parameter(s) supplied"
            + (f", including {', '.join(highlighted_params)}." if highlighted_params else ".")
        )

    if not dependency_items and dependencies_text:
        dependency_items.append(dependencies_text)

    if dependency_items:
        builder.add_heading("3.4 Technical Dependencies", 2)
        builder.add_bullet_list(dependency_items)

    # 4. Message Flow
    print("   [6/10] Message Flow...")
    builder.add_page_break()
    builder.add_heading("4. Message Flow", 1)
    builder.add_paragraph(
        f"A message comes from {source_system} into SAP CI through the configured sender channel."
    )
    builder.add_paragraph(
        f"SAP CI processes the payload and sends the final message to {target_system} through the receiver channel."
    )

    if process_flow_data:
        builder.add_heading("4.1 Process Flow", 2)
        steps = process_flow_data.get("steps", []) if isinstance(process_flow_data.get("steps"), list) else []
        if steps:
            builder.add_table(
                ["Step", "Description"],
                [[str(idx + 1), str(step)] for idx, step in enumerate(steps)]
            )
        summary_rows = dict_to_rows({
            "target_system": process_flow_data.get("target_system", ""),
            "trigger": process_flow_data.get("trigger", ""),
        })
        if summary_rows:
            builder.add_table(["Flow Attribute", "Value"], summary_rows)

    if message_flows:
        builder.add_heading("4.2 Message Flow Connections", 2)
        builder.add_table(
            ["Connection", "Name"],
            [[f"{src} -> {tgt}", name] for src, tgt, name in message_flows],
        )
    elif sequence_flows:
        builder.add_heading("4.2 Sequence Flow Connections", 2)
        builder.add_table(
            ["Connection", "Label"],
            [[f"{src} -> {tgt}", name] for src, tgt, name in sequence_flows],
        )

    if 'integration_flow' in diagram_bytes:
        builder.add_heading("4.3 Integration Flow Diagram", 2)
        builder.add_image(diagram_bytes['integration_flow'], width=6.2, caption="Integration Flow Diagram")

    # 5. Technical Description
    print("   [7/10] Technical Description...")
    builder.add_page_break()
    builder.add_heading("5. Technical Description", 1)

    process_ai = {
        str(item.get("name", "")).strip(): item
        for item in get_list("integration_processes")
        if isinstance(item, dict)
    }

    def render_process_block(
        processes: List[Dict[str, Any]],
        heading: str,
        process_diagrams: Optional[List[Tuple[str, bytes]]] = None,
    ):
        builder.add_heading(heading, 2)
        if not processes:
            builder.add_paragraph("No processes available in this section.", italic=True)
            return

        if process_diagrams:
            for process_name, image_bytes in process_diagrams:
                builder.add_image(image_bytes, width=5.8, caption=f"{process_name} Diagram")

        summary_rows = summarize_processes(processes)
        if summary_rows:
            builder.add_table(["Process Name", "Sequence Flows", "Flow Elements"], summary_rows)

        for idx, process in enumerate(processes, start=1):
            proc_name = process.get("name", f"Process {idx}")
            builder.add_heading(proc_name, 3)

            ai_entry = process_ai.get(proc_name, {})
            description = str(ai_entry.get("description", "")).strip()
            if not description:
                description = fallback_process_description(process)
            if description:
                builder.add_paragraph(description)

            key_activities = ai_entry.get("key_activities", []) if isinstance(ai_entry.get("key_activities"), list) else []
            if not key_activities:
                key_activities = [
                    str(child.get("name", "")).strip()
                    for child in parser.extract_child_properties(process.get("element"))
                    if str(child.get("name", "")).strip()
                    and str(child.get("tag", "")).strip().lower() not in {"startevent", "endevent", "extensionelements"}
                ][:6]
            if key_activities:
                builder.add_table(["Key Activity"], [[str(activity)] for activity in key_activities])

            steps = ai_entry.get("steps", []) if isinstance(ai_entry.get("steps"), list) else []
            if not steps:
                steps = fallback_process_steps(process.get("element"))
            if steps:
                builder.add_table(["Step", "Description"], [[str(i + 1), str(step)] for i, step in enumerate(steps)])

            process_elem = process.get("element")
            if process_elem is None:
                continue

            components = parser.extract_components_from_process(process_elem)
            if components:
                component_rows = filter_display_rows([[key, value] for _, key, value in components if key or value])
                if component_rows:
                    builder.add_table(["Key", "Value"], component_rows)

            child_props = parser.extract_child_properties(process_elem)
            for child in child_props:
                child_heading = child.get("heading", "").strip()
                child_tag = str(child.get("tag", "")).strip().lower()
                child_activity = str(child.get("activity_type", "")).strip().lower()
                props = child.get("properties", [])
                if child_tag in {"startevent", "endevent"}:
                    continue
                if child_activity in {"mapping", "script"}:
                    continue
                props = filter_display_rows(props)
                if child_heading and props:
                    builder.add_heading(f"{child_heading} Properties", 4)
                    builder.add_table(["Key", "Value"], props)

    render_process_block(main_processes, "5.1 Main Integration Process")
    if local_processes:
        render_process_block(local_processes, "5.2 Local Integration Process", local_process_diagrams)

    # Sender
    builder.add_heading("5.3 Sender", 2)
    if 'sender' in diagram_bytes:
        builder.add_image(diagram_bytes['sender'], width=5.2, caption="Sender Configuration")
    sender_ai = get_dict("sender_details")
    sender_description = str(sender_ai.get("description", "")).strip() if sender_ai else ""
    if not sender_description:
        sender_description = build_adapter_description(sender_prop_map, "sender")
    if sender_description:
        builder.add_paragraph(sender_description)
    sender_summary_rows = build_adapter_summary_rows(sender_prop_map, "sender")
    if sender_summary_rows:
        builder.add_table(["Attribute", "Value"], sender_summary_rows)
    else:
        builder.add_paragraph("No sender configuration found.", italic=True)

    # Receiver
    builder.add_heading("5.4 Receiver", 2)
    if 'receiver' in diagram_bytes:
        builder.add_image(diagram_bytes['receiver'], width=5.2, caption="Receiver Configuration")
    receiver_ai = get_dict("receiver_details")
    receiver_description = str(receiver_ai.get("description", "")).strip() if receiver_ai else ""
    if not receiver_description:
        receiver_description = build_adapter_description(receiver_prop_map, "receiver")
    if receiver_description:
        builder.add_paragraph(receiver_description)
    receiver_summary_rows = build_adapter_summary_rows(receiver_prop_map, "receiver")
    if receiver_summary_rows:
        builder.add_table(["Attribute", "Value"], receiver_summary_rows)
    else:
        builder.add_paragraph("No receiver configuration found.", italic=True)

    # Mappings
    builder.add_heading("5.5 Mappings", 2)
    mapping_files_by_name: Dict[str, str] = {}
    for rel_path in list_artifact_paths([".xsl", ".xslt", ".mmap"]):
        mapping_files_by_name[Path(rel_path).stem.lower()] = rel_path

    if mapping_props:
        for idx, mapping in enumerate(mapping_props, start=1):
            mapping_map = props_to_map(mapping)
            mapping_name = (
                mapping_map.get("mappingname")
                or Path(mapping_map.get("mappinguri", "")).stem
                or Path(mapping_map.get("mappingpath", "")).stem
                or f"Mapping {idx}"
            )
            mapping_file = (
                mapping_files_by_name.get(str(mapping_name).lower())
                or mapping_map.get("mappinguri")
                or mapping_map.get("mappingpath")
                or "Not found in project artifacts"
            )
            builder.add_heading(str(mapping_name), 3, collapsed=True)
            for summary_line in build_mapping_summary(str(mapping_name), str(mapping_file)):
                builder.add_paragraph(summary_line)
            source_object, target_object = extract_mapping_relation(str(mapping_name), str(mapping_file))
            mapping_relation_rows = [
                ["Source Message / Object", source_object],
                ["Target Message / Object", target_object],
                ["Mapping File", str(mapping_file)],
            ]
            builder.add_table(["Attribute", "Value"], mapping_relation_rows)
    else:
        builder.add_paragraph("No mapping activities found in this integration flow.", italic=True)

    # Security
    builder.add_heading("5.6 Security", 2)
    security_ai = get_dict("security_config")
    if security_ai:
        builder.add_table(["Security Aspect", "Details"], dict_to_rows(security_ai))
    if security_props:
        filtered_security_rows = filter_display_rows(security_props)
        if filtered_security_rows:
            builder.add_table(["Key", "Value"], filtered_security_rows)
        elif not security_ai:
            builder.add_paragraph("No explicit security properties found.", italic=True)
    else:
        builder.add_paragraph("No explicit security properties found.", italic=True)

    # Groovy Scripts
    builder.add_heading("5.7 Groovy Scripts", 2)
    groovy_ai = get_dict("groovy_scripts")
    if groovy_ai.get("overview"):
        builder.add_paragraph(str(groovy_ai.get("overview")))

    groovy_summaries: Dict[str, Dict[str, Any]] = {}
    groovy_entries = groovy_ai.get("scripts", []) if isinstance(groovy_ai.get("scripts"), list) else []
    for entry in groovy_entries:
        if not isinstance(entry, dict):
            continue
        key = str(entry.get("name", "")).strip().lower()
        if key:
            groovy_summaries[key] = entry

    if groovy_scripts:
        for idx, script in enumerate(groovy_scripts, start=1):
            script_name = str(script.get("file_name") or script.get("name") or f"Script {idx}")
            content = str(script.get("content", "") or "")
            builder.add_heading(script_name, 3, collapsed=True)

            summary_entry = groovy_summaries.get(script_name.lower(), {})
            summary_text = str(summary_entry.get("purpose", "")).strip()
            if not summary_text:
                function_count = len(script.get("functions", [])) if isinstance(script.get("functions"), list) else 0
                imports = script.get("imports", []) if isinstance(script.get("imports"), list) else []
                summary_parts = [f"{script_name} is used within the integration flow."]
                if function_count:
                    summary_parts.append(f"It contains {function_count} function(s).")
                if imports:
                    summary_parts.append(f"It imports {len(imports)} module(s).")
                summary_text = " ".join(summary_parts)
            builder.add_paragraph(summary_text)

            if content.strip():
                builder.add_code_block(content, "Groovy", max_lines=None)
            else:
                builder.add_paragraph("No script content available.", italic=True)
    else:
        builder.add_paragraph("No Groovy scripts in this integration flow.", italic=True)

    # Error handling and logging
    builder.add_heading("5.8 Error Handling & Logging", 2)
    error_ai = get_dict("error_handling")
    if error_ai:
        builder.add_table(["Aspect", "Details"], dict_to_rows(error_ai))

    if 'exception_subprocess' in diagram_bytes:
        builder.add_heading("Exception SubProcess Diagram", 3)
        builder.add_image(
            diagram_bytes['exception_subprocess'],
            width=5.4,
            caption="Exception SubProcess (Copied from Integration Flow Diagram)",
        )

    if exception_props:
        for idx, exc in enumerate(exception_props, start=1):
            builder.add_heading(f"Exception SubProcess {idx} Properties", 3)
            sub_rows = filter_display_rows(exc.get("subproc_props", []))
            if sub_rows:
                builder.add_table(["Key", "Value"], sub_rows)

            for child in exc.get("children", []):
                child_rows = filter_display_rows(child.get("props", []))
                if not child_rows:
                    continue
                child_title = f"Child Element: {child.get('tag', '')} {child.get('name', '')}".strip()
                builder.add_heading(child_title, 4)
                builder.add_table(["Key", "Value"], child_rows)
    else:
        builder.add_paragraph("No exception subprocess configuration found.", italic=True)

    # 6. Version and Metadata
    print("   [8/10] Version and Metadata...")
    builder.add_page_break()
    builder.add_heading("6. Version and Metadata", 1)
    combined_metadata = metadata.copy()
    combined_metadata.update(get_dict("metadata"))
    meta_rows = dict_to_rows(combined_metadata)
    if meta_rows:
        builder.add_table(["Key", "Value"], meta_rows)
    else:
        builder.add_paragraph("No metadata found in the integration flow.", italic=True)

    # 7. Appendix
    print("   [9/10] Appendix...")
    builder.add_page_break()
    builder.add_heading("7. Appendix", 1)
    appendix = get_dict("appendix")

    artifacts = [f"iFlow: {iflow_name}.iflw"]
    artifacts.extend([f"Groovy Script: {s.get('file_name', 'Unknown')}" for s in groovy_scripts])
    artifacts.extend([f"Schema: {s.get('file_name', 'Unknown')}" for s in schemas])
    artifacts.append(f"Message Flows: {len(message_flows)}")
    artifacts.append(f"Sequence Flows: {len(sequence_flows)}")
    artifacts.append(f"Mapping Activities: {len(mapping_props)}")

    appendix_artifacts = appendix.get("artifacts", [])
    if isinstance(appendix_artifacts, list):
        artifacts.extend([str(item) for item in appendix_artifacts])

    builder.add_heading("7.1 Technical Artifacts", 2)
    builder.add_bullet_list(artifacts)

    if schemas:
        builder.add_heading("7.2 Schemas", 2)
        schema_rows = []
        for schema in schemas:
            schema_rows.append([
                str(schema.get("file_name", "")),
                str(schema.get("target_namespace", "")),
                str(len(schema.get("elements", []))) if isinstance(schema.get("elements"), list) else "0",
                str(len(schema.get("complex_types", []))) if isinstance(schema.get("complex_types"), list) else "0",
            ])
        builder.add_table(["File", "Namespace", "Elements", "Complex Types"], schema_rows)

    if parameters:
        builder.add_heading("7.3 Externalized Parameters", 2)
        param_rows = build_parameter_usage_rows(parameters)
        if param_rows:
            builder.add_table(["Parameter", "Value", "Used In"], param_rows)
        else:
            builder.add_paragraph("No externalized parameters found.", italic=True)

    if is_extended_scope:
        glossary = appendix.get("glossary", []) if isinstance(appendix.get("glossary"), list) else []
        if glossary:
            builder.add_heading("7.4 Glossary", 2)
            builder.add_bullet_list([str(item) for item in glossary])

        references = appendix.get("references", "")
        if references:
            builder.add_heading("7.5 References", 2)
            builder.add_paragraph(str(references))

        if ai_stats:
            builder.add_heading("7.6 Generation Statistics", 2)
            stat_rows = [
                ["API Calls", str(ai_stats.get("api_calls", 0))],
                ["Batch Calls", str(ai_stats.get("batch_calls", 0))],
                ["Cache Hits", str(ai_stats.get("cache_hits", 0))],
                ["Cache Hit Rate", f"{ai_stats.get('cache_hit_rate', 0)}%"],
                ["Failures", str(ai_stats.get("failures", 0))],
            ]
            builder.add_table(["Statistic", "Value"], stat_rows)

        if file_type_summary:
            builder.add_heading("7.7 File Type Summary", 2)
            type_rows = [[str(ext), str(count)] for ext, count in file_type_summary.items()]
            builder.add_table(["File Type", "Count"], type_rows)

        if all_files:
            builder.add_heading("7.8 File Inventory", 2)
            inventory_rows: List[List[str]] = []
            max_inventory_rows = 160
            sorted_files = sorted(
                all_files,
                key=lambda item: str(item.get("relative_path", "")).lower(),
            )
            for entry in sorted_files[:max_inventory_rows]:
                inventory_rows.append(
                    [
                        str(entry.get("relative_path", "")),
                        str(entry.get("category", "")),
                        str(entry.get("extension", "")),
                        str(entry.get("size_bytes", "")),
                    ]
                )

            builder.add_table(["File", "Category", "Type", "Size (bytes)"], inventory_rows)
            if len(sorted_files) > max_inventory_rows:
                builder.add_paragraph(
                    f"Inventory truncated in document view: showing {max_inventory_rows} of {len(sorted_files)} files.",
                    italic=True,
                )

        if artifact_coverage or text_artifacts:
            builder.add_heading("7.9 Artifact Evidence", 2)

            if artifact_coverage:
                builder.add_table(
                    ["Artifact Coverage Aspect", "Details"],
                    dict_to_rows(
                        {
                            "analyzed_file_types": artifact_coverage.get("analyzed_file_types", []),
                            "critical_non_iflow_artifacts": artifact_coverage.get(
                                "critical_non_iflow_artifacts", []
                            ),
                            "observations": artifact_coverage.get("observations", []),
                        }
                    ),
                )

            if text_artifacts:
                evidence_rows: List[List[str]] = []
                for artifact in text_artifacts[:60]:
                    signals = artifact.get("signal_lines", [])
                    signal_text = " | ".join([str(line) for line in signals[:2]]) if isinstance(signals, list) else ""
                    evidence_rows.append(
                        [
                            str(artifact.get("relative_path", "")),
                            str(artifact.get("category", "")),
                            signal_text,
                        ]
                    )

                builder.add_table(["File", "Category", "Evidence Snippets"], evidence_rows)

    print("   [10/10] Finalizing...")
    output_path = builder.save()

    print("\n" + "=" * 70)
    print(f"  [OK] Document generated: {output_path}")
    print("=" * 70)
    return output_path


# Backward compatibility for older imports
build_specification = build_specification_document
