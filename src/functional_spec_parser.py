"""Utilities to extract optional functional specification context for AI prompts."""

from __future__ import annotations

import logging
import os
import re
from pathlib import Path
from typing import Any, Dict, List, Optional

try:
    from docx import Document as DocxDocument
except Exception:  # pragma: no cover - optional at runtime
    DocxDocument = None


SUPPORTED_SPEC_EXTENSIONS = {".docx", ".doc", ".txt", ".md", ".rtf"}

EXCLUDED_AUTO_SPEC_DIR_NAMES = {
    ".git",
    ".idea",
    ".vscode",
    "__pycache__",
    "node_modules",
    "output",
    "temp",
    "venv",
}

STOP_WORDS = {
    "the",
    "and",
    "for",
    "that",
    "with",
    "from",
    "this",
    "shall",
    "will",
    "into",
    "must",
    "should",
    "have",
    "has",
    "are",
    "was",
    "were",
    "via",
    "per",
    "using",
    "integration",
    "process",
    "business",
    "functional",
    "specification",
    "hyperlink",
    "http",
    "https",
    "xsd",
    "xmlns",
    "appinfo",
    "element",
    "string",
    "name",
    "type",
}

SECTION_KEYWORDS = {
    "overview",
    "background",
    "scope",
    "objective",
    "purpose",
    "business requirement",
    "business requirements",
    "functional requirement",
    "interface",
    "flow",
    "scenario",
    "assumption",
    "dependency",
    "integration",
    "mapping",
    "sender",
    "receiver",
    "error",
    "exception",
    "security",
    "performance",
    "volume",
    "sla",
    "test",
    "contact",
    "reconciliation",
    "change management",
}

NOISE_LINE_PATTERNS = [
    re.compile(pattern, re.IGNORECASE)
    for pattern in [
        r"\bFORMCHECKBOX\b",
        r"\bSummaryInformation\b",
        r"\bDocumentSummaryInformation\b",
        r"\bVisioInformation\b",
        r"\bTableStyle\b",
        r"\bDefault Paragraph Font\b",
        r"\bTimes New Roman\b",
        r"\bMICROSOFT SANS SERIF\b",
        r"\bGeneral Everyday Printing\b",
        r"\bAdobe Photoshop\b",
        r"\bOffice\\\d+\\\w+\.EXE\b",
        r"\bPCL\b",
        r"^#\([\d;@A-Za-z]+\)$",
        r"^\\\\[^\s]+",
        r"^HYPERLINK\b",
        r"^https?://[^\s]+$",
        r"<\s*/?xsd:",
        r"xmlns:xsd=",
        r"\b(?:xsd|xmlns|appinfo|complextype|targetnamespace|schemaid|sequencenum)\b",
        r"<\s*/?w:",
        r"<\s*/?o:",
        r"<\?xml",
    ]
]


def _dedupe_keep_order(items: List[str], max_items: int = 12) -> List[str]:
    """Deduplicate non-empty strings preserving first occurrence order."""
    result: List[str] = []
    seen: set[str] = set()
    for item in items:
        clean = _normalize_text(item)
        if not clean:
            continue
        key = clean.lower()
        if key in seen:
            continue
        seen.add(key)
        result.append(clean)
        if len(result) >= max_items:
            break
    return result


def _extract_lines_by_keywords(
    lines: List[str],
    keywords: List[str],
    max_items: int = 10,
) -> List[str]:
    """Extract lines containing any keyword while keeping order and uniqueness."""
    pattern = re.compile("|".join([re.escape(word) for word in keywords]), re.IGNORECASE)
    selected: List[str] = []
    for line in lines:
        if len(line) < 8:
            continue
        if pattern.search(line):
            selected.append(line)
    return _dedupe_keep_order(selected, max_items=max_items)


def _extract_step_candidates(lines: List[str], max_items: int = 12) -> List[str]:
    """Extract ordered or bullet-like process steps from functional text."""
    step_like: List[str] = []
    for line in lines:
        normalized = _normalize_text(line)
        if not normalized:
            continue
        if re.match(r"^(\d+|[ivxlcdm]+)[\.)]\s+", normalized, flags=re.IGNORECASE):
            step_like.append(normalized)
            continue
        if re.match(r"^[-*]\s+", normalized):
            step_like.append(normalized)
    return _dedupe_keep_order(step_like, max_items=max_items)


def _extract_top_terms(text: str, max_items: int = 12) -> List[str]:
    """Extract frequent non-trivial terms as thematic hints."""
    tokens = re.findall(r"[A-Za-z][A-Za-z0-9_\-]{2,}", text.lower())
    freq: Dict[str, int] = {}
    for token in tokens:
        if token in STOP_WORDS:
            continue
        if token.isdigit():
            continue
        freq[token] = freq.get(token, 0) + 1

    ranked = sorted(freq.items(), key=lambda item: (-item[1], item[0]))
    return [term for term, _ in ranked[:max_items]]


def _is_likely_noise_line(line: str) -> bool:
    """Heuristic filter for formatting artifacts and embedded technical noise."""
    candidate = _normalize_text(line)
    if not candidate:
        return True

    if len(candidate) > 450:
        return True

    for pattern in NOISE_LINE_PATTERNS:
        if pattern.search(candidate):
            return True

    alpha_chars = sum(1 for ch in candidate if ch.isalpha())
    symbol_chars = sum(1 for ch in candidate if not ch.isalnum() and not ch.isspace())
    if alpha_chars == 0 and symbol_chars > 4:
        return True

    if alpha_chars > 0 and (symbol_chars / max(len(candidate), 1)) > 0.45:
        return True

    # Exclude dense XML-like lines that are rarely meaningful in functional narratives.
    if candidate.count("<") >= 3 and candidate.count(">") >= 3:
        return True

    return False


def _sanitize_extracted_text(text: str) -> str:
    """Remove obvious extraction artifacts while preserving business narrative lines."""
    lines = [_normalize_text(line) for line in text.splitlines()]
    cleaned: List[str] = []
    seen: set[str] = set()

    for line in lines:
        if len(line) < 3:
            continue
        if _is_likely_noise_line(line):
            continue

        key = line.lower()
        if key in seen:
            continue
        seen.add(key)
        cleaned.append(line)

    return "\n".join(cleaned).strip()


def _is_heading_line(line: str) -> bool:
    """Estimate whether a line behaves like a section heading."""
    text = _normalize_text(line)
    if len(text) < 4 or len(text) > 120:
        return False

    if text.endswith(":"):
        text = text[:-1].strip()

    if _is_likely_noise_line(text):
        return False

    lower = text.lower()
    if any(keyword in lower for keyword in SECTION_KEYWORDS):
        return True

    if re.match(r"^\d+(?:\.\d+)*\s+[A-Za-z].*", text):
        return True

    return False


def _extract_section_map(lines: List[str], max_sections: int = 24) -> List[Dict[str, Any]]:
    """Build a dynamic section map from heading-like lines and following evidence."""
    sections: List[Dict[str, Any]] = []
    current: Dict[str, Any] | None = None

    for line in lines:
        normalized = _normalize_text(line)
        if not normalized:
            continue

        if _is_heading_line(normalized):
            if current and current.get("content"):
                sections.append(current)

            current = {"title": normalized, "content": []}
            continue

        if current is None:
            continue

        content: List[str] = current.get("content", [])
        if len(content) < 12:
            content.append(normalized)
            current["content"] = content

    if current and current.get("content"):
        sections.append(current)

    # Deduplicate by title, keep first occurrence.
    deduped: List[Dict[str, Any]] = []
    seen_titles: set[str] = set()
    for section in sections:
        title = str(section.get("title", "")).strip()
        if not title:
            continue

        key = title.lower()
        if key in seen_titles:
            continue
        seen_titles.add(key)

        deduped.append(section)
        if len(deduped) >= max_sections:
            break

    return deduped


def _build_llm_context_from_analysis(analysis: Dict[str, Any], max_chars: int = 12000) -> str:
    """Create compact, business-oriented context block for LLM consumption."""
    lines: List[str] = []

    lines.append("Functional specification synthesized context")
    lines.append(f"Documents analyzed: {analysis.get('document_count', 0)}")

    top_terms = analysis.get("top_terms", [])
    if isinstance(top_terms, list) and top_terms:
        lines.append("Top terms: " + ", ".join([str(term) for term in top_terms[:10]]))

    for key, label in [
        ("business_requirements", "Business requirements"),
        ("interface_points", "Interface points"),
        ("assumptions", "Assumptions"),
        ("constraints_and_risks", "Constraints and risks"),
        ("operational_indicators", "Operational indicators"),
        ("process_steps", "Process steps"),
        ("system_mentions", "System mentions"),
    ]:
        values = analysis.get(key, [])
        if isinstance(values, list) and values:
            lines.append(f"{label}:")
            for item in values[:12]:
                lines.append(f"- {item}")

    section_map = analysis.get("section_map", [])
    if isinstance(section_map, list) and section_map:
        lines.append("Section evidence:")
        for section in section_map[:14]:
            title = str(section.get("title", "")).strip()
            if not title:
                continue
            lines.append(f"- Section: {title}")

            content = section.get("content", [])
            if isinstance(content, list):
                for snippet in content[:3]:
                    lines.append(f"  * {snippet}")

    context = "\n".join(lines).strip()
    if len(context) <= max_chars:
        return context

    return context[:max_chars].rstrip() + "\n[Functional specification synthesized context truncated]"


def _build_functional_spec_analysis(blocks: List[Dict[str, str]]) -> Dict[str, Any]:
    """Build structured analysis from extracted functional-spec text blocks."""
    if not blocks:
        return {}

    all_lines: List[str] = []
    source_files: List[str] = []
    combined_parts: List[str] = []

    for block in blocks:
        source = str(block.get("source", "")).strip()
        text = _sanitize_extracted_text(str(block.get("text", "")))
        if not text:
            continue
        if source:
            source_files.append(source)
        combined_parts.append(text)
        all_lines.extend([_normalize_text(line) for line in text.splitlines() if _normalize_text(line)])

    if not combined_parts:
        return {}

    combined_text = "\n".join(combined_parts)

    section_map = _extract_section_map(all_lines)

    requirements = _extract_lines_by_keywords(
        all_lines,
        ["require", "must", "shall", "need", "mandatory", "acceptance", "scope"],
        max_items=14,
    )
    interface_points = _extract_lines_by_keywords(
        all_lines,
        ["interface", "api", "endpoint", "sender", "receiver", "source", "target", "protocol"],
        max_items=12,
    )
    assumptions = _extract_lines_by_keywords(
        all_lines,
        ["assumption", "depends", "dependency", "precondition", "expected", "availability"],
        max_items=10,
    )
    constraints = _extract_lines_by_keywords(
        all_lines,
        ["constraint", "limitation", "risk", "fallback", "exception", "error", "retry"],
        max_items=10,
    )
    operational = _extract_lines_by_keywords(
        all_lines,
        ["frequency", "schedule", "daily", "hourly", "volume", "throughput", "sla"],
        max_items=10,
    )
    step_candidates = _extract_step_candidates(all_lines, max_items=14)

    systems = _extract_lines_by_keywords(
        all_lines,
        ["system", "application", "erp", "sap", "service", "platform"],
        max_items=12,
    )

    top_terms = _extract_top_terms(combined_text, max_items=12)

    analysis: Dict[str, Any] = {
        "source_files": _dedupe_keep_order(source_files, max_items=30),
        "document_count": len(_dedupe_keep_order(source_files, max_items=1000)),
        "top_terms": top_terms,
        "business_requirements": requirements,
        "interface_points": interface_points,
        "assumptions": assumptions,
        "constraints_and_risks": constraints,
        "operational_indicators": operational,
        "process_steps": step_candidates,
        "system_mentions": systems,
        "section_map": section_map,
    }

    summary_parts: List[str] = []
    if requirements:
        summary_parts.append(f"requirements={len(requirements)}")
    if interface_points:
        summary_parts.append(f"interfaces={len(interface_points)}")
    if step_candidates:
        summary_parts.append(f"steps={len(step_candidates)}")
    if operational:
        summary_parts.append(f"operational-signals={len(operational)}")

    if summary_parts:
        analysis["analysis_summary"] = "Functional-spec signals detected: " + ", ".join(summary_parts)
    else:
        analysis["analysis_summary"] = "Functional specification loaded, but explicit requirement/interface markers were limited."

    analysis["llm_context"] = _build_llm_context_from_analysis(analysis)

    return analysis


def _is_excluded_dir_name(name: str) -> bool:
    """Check whether a directory should be excluded from auto-discovery."""
    return name.lower() in EXCLUDED_AUTO_SPEC_DIR_NAMES


def _iter_supported_files(root: Path, max_depth: Optional[int] = None) -> List[Path]:
    """Walk root and collect supported files while pruning excluded directories."""
    if not root.exists() or not root.is_dir():
        return []

    files: List[Path] = []
    for current_root, dirnames, filenames in os.walk(root):
        try:
            rel_parts = Path(current_root).relative_to(root).parts
            depth = len(rel_parts)
        except Exception:
            depth = 0

        if max_depth is not None and depth >= max_depth:
            dirnames[:] = []

        dirnames[:] = [d for d in dirnames if not _is_excluded_dir_name(d)]
        for file_name in filenames:
            suffix = Path(file_name).suffix.lower()
            if suffix in SUPPORTED_SPEC_EXTENSIONS:
                files.append(Path(current_root) / file_name)
    return files


def _score_candidate(file_path: Path, base_dir: Path, parent_dir: Path) -> int:
    """Score a potential functional specification file path."""
    stem = file_path.stem.lower()
    full_name = file_path.name.lower()

    score = 0

    # Baseline preference for Word documents.
    if file_path.suffix.lower() in {".doc", ".docx"}:
        score += 2
    else:
        score += 1

    # Positive keyword signals.
    if "functional" in stem:
        score += 8
    if "specification" in stem:
        score += 6
    if re.search(r"\bspec\b", stem):
        score += 4
    if re.search(r"\brequirement\b", stem):
        score += 3
    if "business" in stem:
        score += 2

    # Negative signals for generated or irrelevant docs.
    if "techspec" in stem or "technical specification" in stem:
        score -= 10
    if re.search(r"technical[_\- ]*specification", stem):
        score -= 10
    if "readme" in stem:
        score -= 6
    if "requirements" in stem:
        score -= 8
    if "changelog" in stem:
        score -= 4

    path_parts = [part.lower() for part in file_path.parts]
    if any(part in EXCLUDED_AUTO_SPEC_DIR_NAMES for part in path_parts):
        score -= 12

    # Prefer files close to the integration project directory.
    if file_path.parent == base_dir:
        score += 4
    elif file_path.parent == parent_dir:
        score += 3

    # Small depth penalty from root for broad directories.
    try:
        depth = len(file_path.relative_to(base_dir).parts)
        score -= max(depth - 4, 0)
    except Exception:
        pass

    # Prefer clearer keyword combinations.
    if "functional" in full_name and ("spec" in full_name or "specification" in full_name):
        score += 4

    return score


def discover_functional_spec_path(
    input_path: Path,
    logger: Optional[logging.Logger] = None,
) -> Optional[Path]:
    """
    Auto-detect the best functional specification file near the input project.

    Search strategy:
    - input directory (or ZIP parent directory)
    - parent directory of that base
    """
    log = logger or logging.getLogger(__name__)
    input_path = Path(input_path)

    base_dir = input_path if input_path.is_dir() else input_path.parent
    parent_dir = base_dir.parent

    search_roots: List[tuple[Path, Optional[int]]] = []
    for root, max_depth in [(base_dir, 6), (parent_dir, 1)]:
        if root and root.exists() and root.is_dir():
            if all(existing_root != root for existing_root, _ in search_roots):
                search_roots.append((root, max_depth))

    candidates: List[tuple[int, Path]] = []
    seen: set[Path] = set()

    for root, max_depth in search_roots:
        for file_path in _iter_supported_files(root, max_depth=max_depth):
            resolved = file_path.resolve()
            if resolved in seen:
                continue
            seen.add(resolved)

            score = _score_candidate(file_path, base_dir, parent_dir)
            if score >= 7:
                candidates.append((score, file_path))

    if not candidates:
        return None

    candidates.sort(key=lambda item: (-item[0], len(str(item[1]))))
    best_score, best_path = candidates[0]
    log.debug(f"Auto-selected functional spec candidate: {best_path} (score={best_score})")
    return best_path


def _normalize_text(text: str) -> str:
    """Normalize whitespace while preserving paragraph-like newlines."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _read_text_with_fallback(path: Path) -> str:
    """Read text content with conservative encoding fallbacks."""
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except Exception:
            continue

    return path.read_text(encoding="utf-8", errors="ignore")


def _extract_docx_text(path: Path) -> str:
    """Extract paragraphs and table text from DOCX documents."""
    if DocxDocument is None:
        raise RuntimeError("python-docx is not available")

    doc = DocxDocument(str(path))
    parts: List[str] = []

    for paragraph in doc.paragraphs:
        text = _normalize_text(paragraph.text)
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [_normalize_text(cell.text) for cell in row.cells]
            cells = [c for c in cells if c]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _extract_rtf_text(path: Path) -> str:
    """Best-effort extraction for RTF content without extra dependencies."""
    text = _read_text_with_fallback(path)
    # Convert common escaped hex entities first.
    text = re.sub(r"\\'[0-9a-fA-F]{2}", " ", text)
    # Treat paragraph markers as line breaks.
    text = re.sub(r"\\par[d]?", "\n", text)
    # Strip remaining control words.
    text = re.sub(r"\\[a-zA-Z]+-?\d* ?", " ", text)
    # Remove grouping braces.
    text = text.replace("{", " ").replace("}", " ")
    return _normalize_text(text)


def _extract_doc_text_with_word(doc_path: Path, logger: logging.Logger) -> str:
    """Try reading legacy .doc via local Word automation on Windows."""
    try:
        from win32com.client import DispatchEx  # type: ignore
    except Exception:
        return ""

    app = None
    document = None

    try:
        app = DispatchEx("Word.Application")
        app.Visible = False
        app.DisplayAlerts = 0
        document = app.Documents.Open(
            str(doc_path.resolve()),
            ConfirmConversions=False,
            ReadOnly=True,
            AddToRecentFiles=False,
        )
        content = document.Content.Text or ""
        return _normalize_text(content)
    except Exception as exc:
        logger.debug(f"Word automation failed for {doc_path.name}: {exc}")
        return ""
    finally:
        if document is not None:
            try:
                document.Close(False)
            except Exception:
                pass
        if app is not None:
            try:
                app.Quit()
            except Exception:
                pass


def _extract_doc_text_heuristic(path: Path) -> str:
    """Fallback extraction for .doc binary content using printable string heuristics."""
    raw = path.read_bytes()

    utf16_chunks = re.findall(rb"(?:[\x20-\x7E]\x00){14,}", raw)
    latin_chunks = re.findall(rb"[\x20-\x7E]{28,}", raw)

    candidates: List[str] = []
    for chunk in utf16_chunks:
        decoded = chunk.decode("utf-16le", errors="ignore")
        if decoded:
            candidates.append(decoded)

    for chunk in latin_chunks:
        decoded = chunk.decode("latin-1", errors="ignore")
        if decoded:
            candidates.append(decoded)

    lines: List[str] = []
    seen: set[str] = set()

    for candidate in candidates:
        for raw_line in candidate.splitlines():
            line = _normalize_text(raw_line)
            if len(line) < 8:
                continue
            if re.search(r"[A-Za-z]{3}", line) is None:
                continue

            key = line.lower()
            if key in seen:
                continue
            seen.add(key)
            lines.append(line)

    return "\n".join(lines)


def _extract_doc_text(path: Path, logger: logging.Logger) -> str:
    """Extract text from legacy .doc using the best available method."""
    text = _extract_doc_text_with_word(path, logger)
    if text:
        return text

    return _extract_doc_text_heuristic(path)


def _extract_file_text(path: Path, logger: logging.Logger) -> str:
    """Extract text based on file extension."""
    suffix = path.suffix.lower()

    if suffix == ".docx":
        return _extract_docx_text(path)
    if suffix == ".doc":
        return _extract_doc_text(path, logger)
    if suffix == ".rtf":
        return _extract_rtf_text(path)
    if suffix in {".txt", ".md"}:
        return _read_text_with_fallback(path)

    return ""


def load_functional_spec_context(
    spec_path: Optional[Path],
    max_chars: int = 15000,
    logger: Optional[logging.Logger] = None,
) -> Dict[str, Any]:
    """
    Load optional functional-spec content for AI context enrichment.

    Returns a dictionary with keys:
    - context: concatenated text
    - loaded_files: files successfully read
    - ignored_files: unsupported or unreadable files
    - warnings: non-fatal warnings
    - truncated: whether context was truncated to max_chars
    """
    log = logger or logging.getLogger(__name__)

    result: Dict[str, Any] = {
        "context": "",
        "llm_context": "",
        "loaded_files": [],
        "ignored_files": [],
        "warnings": [],
        "truncated": False,
        "analysis": {},
    }

    if not spec_path:
        return result

    spec_path = Path(spec_path)
    if not spec_path.exists():
        result["warnings"].append(f"Functional spec path not found: {spec_path}")
        return result

    candidates: List[Path] = []
    if spec_path.is_file():
        candidates = [spec_path]
    elif spec_path.is_dir():
        for ext in sorted(SUPPORTED_SPEC_EXTENSIONS):
            candidates.extend(spec_path.rglob(f"*{ext}"))
        candidates = sorted(candidates, key=lambda p: str(p).lower())
    else:
        result["warnings"].append(f"Functional spec path is not file/directory: {spec_path}")
        return result

    if not candidates:
        result["warnings"].append("No supported functional specification files found.")
        return result

    remaining = max(max_chars, 0)
    blocks: List[str] = []
    analysis_blocks: List[Dict[str, str]] = []

    for file_path in candidates:
        if remaining <= 0:
            result["truncated"] = True
            break

        if file_path.suffix.lower() not in SUPPORTED_SPEC_EXTENSIONS:
            result["ignored_files"].append(str(file_path))
            continue

        try:
            extracted = _sanitize_extracted_text(_extract_file_text(file_path, log))
        except Exception as exc:
            log.debug(f"Failed to read functional spec file {file_path}: {exc}")
            extracted = ""

        if not extracted:
            result["ignored_files"].append(str(file_path))
            continue

        block = f"Functional specification source: {file_path.name}\n{extracted}"
        if len(block) > remaining:
            block = block[:remaining].rstrip()
            result["truncated"] = True

        blocks.append(block)
        result["loaded_files"].append(str(file_path))
        analysis_blocks.append({"source": file_path.name, "text": extracted})
        remaining -= len(block)

    if blocks:
        result["context"] = "\n\n---\n\n".join(blocks)
        result["analysis"] = _build_functional_spec_analysis(analysis_blocks)
        analysis = result.get("analysis", {})
        if isinstance(analysis, dict):
            llm_context = analysis.get("llm_context", "")
            if isinstance(llm_context, str):
                result["llm_context"] = llm_context
    else:
        result["warnings"].append("No readable functional specification content could be extracted.")

    if result["truncated"]:
        result["warnings"].append(
            f"Functional specification context truncated to {max_chars} characters."
        )
    return result
